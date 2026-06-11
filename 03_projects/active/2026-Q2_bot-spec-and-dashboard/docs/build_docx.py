"""
Build the consolidated bot + dashboard spec DOCX from the source markdown files.

Run:
    python build_docx.py

Output:
    spec.docx (in this same docs/ folder)
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ============================================================
# Configuration
# ============================================================

PROJECT_ROOT = Path(__file__).resolve().parent.parent
BOT_SPEC_DIR = PROJECT_ROOT / "bot-spec"
DASHBOARD_SPEC_DIR = PROJECT_ROOT / "dashboard-spec"
OUTPUT_PATH = Path(__file__).resolve().parent / "spec.docx"

# Brand colors
NAVY_900 = RGBColor(0x0D, 0x18, 0x28)
NAVY_700 = RGBColor(0x22, 0x35, 0x5C)
NAVY_500 = RGBColor(0x3B, 0x5A, 0x9C)
TEAL_600 = RGBColor(0x0D, 0x94, 0x88)
GRAY_700 = RGBColor(0x37, 0x41, 0x51)
GRAY_500 = RGBColor(0x6B, 0x72, 0x80)
GRAY_200 = RGBColor(0xE5, 0xE7, 0xEB)
GRAY_100 = RGBColor(0xF3, 0xF4, 0xF6)
CODE_BG = RGBColor(0xF6, 0xF8, 0xFA)
CODE_FG = RGBColor(0x24, 0x29, 0x2E)

# Fonts
FONT_HEBREW = "Heebo"
FONT_LATIN = "Calibri"
FONT_MONO = "Consolas"


# ============================================================
# Document structure plan
# ============================================================

PLAN = [
    {
        "part": "חלק I — מסגרת ועקרונות חוצים",
        "part_en": "Part I — Framework & Cross-Cutting",
        "files": [
            ("מסגרת ועקרונות הבוט", "00_overview.md"),
            ("חוקים חוצי-תרחיש", "09_cross_cutting.md"),
            ("אינטגרציית CRM ומודל חמימות-ליד", "10_crm_integration.md"),
        ],
        "dir": BOT_SPEC_DIR,
    },
    {
        "part": "חלק II — תרחישי C2B (יוצאים)",
        "part_en": "Part II — C2B Outbound Scenarios",
        "files": [
            ("C1 — Cold Outreach", "01_c1_cold_outreach.md"),
            ("C2 — Trademobile Active w/ Fines", "02_c2_trademobile_active.md"),
            ("C3 — Past Customer Winback", "03_c3_past_customer_winback.md"),
            ("C4 — Trademobile Welcome", "04_c4_trademobile_welcome.md"),
            ("C5 — Free-Tier User Got a Fine", "05_c5_free_user_got_fine.md"),
            ("C6 — Pre-Expiry Retention", "06_c6_pre_expiry_retention.md"),
            ("C7 — Dirty & Quick Winback", "07_c7_dirty_quick_winback.md"),
        ],
        "dir": BOT_SPEC_DIR,
    },
    {
        "part": "חלק III — תרחיש B2C (נכנס)",
        "part_en": "Part III — B2C Inbound",
        "files": [
            ("B1 — Inbound Entry", "08_b1_inbound_entry.md"),
        ],
        "dir": BOT_SPEC_DIR,
    },
    {
        "part": "חלק IV — אפיון לוח הבקרה",
        "part_en": "Part IV — Dashboard Specification",
        "files": [
            ("מילון מדדים", "01_metrics_glossary.md"),
            ("סכמת אירועים", "02_event_schema.md"),
            ("מודל המשפך", "03_funnel_model.md"),
            ("מסכי הדשבורד", "04_screens.md"),
        ],
        "dir": DASHBOARD_SPEC_DIR,
    },
]


# ============================================================
# Helpers — RTL, fonts, formatting
# ============================================================

HEBREW_RANGE = re.compile(r"[֐-׿]")


def has_significant_hebrew(text: str) -> bool:
    """Return True if the text contains >=20% Hebrew characters.
    Used at the paragraph level to decide RTL bidi layout."""
    if not text:
        return False
    letters = [c for c in text if c.isalpha()]
    if not letters:
        return False
    hebrew = [c for c in letters if HEBREW_RANGE.match(c)]
    return len(hebrew) / len(letters) >= 0.20


def set_paragraph_rtl(paragraph, rtl: bool = True):
    """Apply bidi RTL to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    # Remove any existing bidi
    for existing in pPr.findall(qn("w:bidi")):
        pPr.remove(existing)
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1" if rtl else "0")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT


def set_run_font(run, size_pt=11, bold=False, italic=False, color=None,
                 mono=False, hebrew_hint=False):
    """Apply font styling to a run, including East Asian/complex-script font."""
    # Auto-detect Hebrew at the RUN level. Any Hebrew character at all triggers
    # RTL markup — even a single Hebrew word inside an English paragraph.
    # This catches the "mixed bold Hebrew chunk in English prose" case.
    run_text = run.text if run.text else ""
    has_any_hebrew = any('֐' <= c <= '׿' for c in run_text)
    hebrew_hint = hebrew_hint or has_any_hebrew

    font = run.font
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    if color is not None:
        font.color.rgb = color
    if mono:
        font.name = FONT_MONO
        # Also set the complex-script font (used for Hebrew in monospace contexts)
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), FONT_MONO)
        rFonts.set(qn("w:hAnsi"), FONT_MONO)
        rFonts.set(qn("w:cs"), FONT_MONO)
    else:
        # Default sans-serif. For Hebrew text, use Heebo for complex-script.
        font.name = FONT_LATIN
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), FONT_LATIN)
        rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        rFonts.set(qn("w:cs"), FONT_HEBREW)
    # For Hebrew-dominant runs, also set bold/italic on complex-script side,
    # AND mark the run as RTL so Word renders characters in the correct order.
    if hebrew_hint:
        rPr = run._r.get_or_add_rPr()
        # CRITICAL: without <w:rtl/>, Word renders Hebrew characters in LTR
        # order even when the paragraph is bidi. This is the fix for
        # "Hebrew letters appear reversed" symptoms.
        if rPr.find(qn("w:rtl")) is None:
            rPr.append(OxmlElement("w:rtl"))
        lang = rPr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            rPr.append(lang)
        lang.set(qn("w:bidi"), "he-IL")
        if bold:
            bCs = rPr.find(qn("w:bCs"))
            if bCs is None:
                bCs = OxmlElement("w:bCs")
                rPr.append(bCs)
        if italic:
            iCs = rPr.find(qn("w:iCs"))
            if iCs is None:
                iCs = OxmlElement("w:iCs")
                rPr.append(iCs)
        size = rPr.find(qn("w:szCs"))
        if size is None:
            size = OxmlElement("w:szCs")
            rPr.append(size)
        size.set(qn("w:val"), str(size_pt * 2))


def shade_paragraph(paragraph, color: RGBColor):
    """Add background shading to a paragraph (used for code blocks, callouts)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{color}")
    pPr.append(shd)


def add_paragraph_border(paragraph, position="left", color="0D9488", size=24):
    """Add a left border to a paragraph for callout-style styling."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    border = OxmlElement(f"w:{position}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "8")
    border.set(qn("w:color"), color)
    pBdr.append(border)


# ============================================================
# Inline markdown parser — handles bold/italic/code in a single line
# ============================================================

INLINE_RE = re.compile(
    r"(\*\*([^*]+?)\*\*"      # bold
    r"|\*([^*]+?)\*"           # italic (single asterisk)
    r"|`([^`]+?)`"             # inline code
    r"|\[([^\]]+?)\]\(([^)]+?)\))"  # links [text](url)
)


def add_inline_runs(paragraph, text, base_size=11, hebrew_hint=False):
    """Parse inline markdown and add runs to the paragraph."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        # text before the match
        if m.start() > pos:
            chunk = text[pos:m.start()]
            if chunk:
                run = paragraph.add_run(chunk)
                set_run_font(run, size_pt=base_size, hebrew_hint=hebrew_hint or has_significant_hebrew(chunk))

        if m.group(2) is not None:  # bold
            run = paragraph.add_run(m.group(2))
            set_run_font(run, size_pt=base_size, bold=True, hebrew_hint=hebrew_hint or has_significant_hebrew(m.group(2)))
        elif m.group(3) is not None:  # italic
            run = paragraph.add_run(m.group(3))
            set_run_font(run, size_pt=base_size, italic=True, hebrew_hint=hebrew_hint or has_significant_hebrew(m.group(3)))
        elif m.group(4) is not None:  # inline code
            code_text = m.group(4)
            # If the inline code contains Hebrew, don't use mono (Consolas
            # has no Hebrew glyphs). Render as italic+teal — visually
            # distinguishable as a "term" without breaking Hebrew rendering.
            if any('֐' <= c <= '׿' for c in code_text):
                run = paragraph.add_run(code_text)
                set_run_font(run, size_pt=base_size, italic=True,
                             color=TEAL_600, hebrew_hint=True)
            else:
                run = paragraph.add_run(code_text)
                set_run_font(run, size_pt=base_size - 1, mono=True, color=NAVY_700)
        elif m.group(5) is not None:  # link [text](url)
            run = paragraph.add_run(m.group(5))
            set_run_font(run, size_pt=base_size, color=TEAL_600, hebrew_hint=hebrew_hint or has_significant_hebrew(m.group(5)))
            run.underline = True

        pos = m.end()

    if pos < len(text):
        chunk = text[pos:]
        if chunk:
            run = paragraph.add_run(chunk)
            set_run_font(run, size_pt=base_size, hebrew_hint=hebrew_hint or has_significant_hebrew(chunk))


# ============================================================
# Document setup
# ============================================================

def setup_document():
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = FONT_LATIN
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.35

    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style.element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_LATIN)
    rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    rFonts.set(qn("w:cs"), FONT_HEBREW)

    # Heading styles — colors, sizes, weights
    # Level 1 = Part (top of hierarchy). Level 2 = Chapter. Level 3+ = sections.
    heading_sizes = {1: 28, 2: 22, 3: 16, 4: 13, 5: 11.5, 6: 11}
    heading_colors = {1: NAVY_900, 2: NAVY_900, 3: NAVY_700, 4: NAVY_700, 5: NAVY_500, 6: NAVY_500}
    heading_spacing_before = {1: 30, 2: 22, 3: 16, 4: 12, 5: 8, 6: 6}
    heading_spacing_after = {1: 14, 2: 10, 3: 7, 4: 5, 5: 4, 6: 2}

    for level in range(1, 7):
        s = doc.styles[f"Heading {level}"]
        s.font.name = FONT_LATIN
        s.font.size = Pt(heading_sizes[level])
        s.font.bold = True
        s.font.color.rgb = heading_colors[level]
        s.paragraph_format.space_before = Pt(heading_spacing_before[level])
        s.paragraph_format.space_after = Pt(heading_spacing_after[level])
        s.paragraph_format.keep_with_next = True

        rPr = s.element.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            s.element.insert(0, rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), FONT_LATIN)
        rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        rFonts.set(qn("w:cs"), FONT_HEBREW)

    return doc


# ============================================================
# Title page + TOC
# ============================================================

def add_title_page(doc):
    # Vertical spacing before title
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Main title (Hebrew)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rtl(p, True)
    run = p.add_run("Road Protect")
    set_run_font(run, size_pt=36, bold=True, color=NAVY_900)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rtl(p, True)
    run = p.add_run("אפיון בוט WhatsApp ולוח בקרה")
    set_run_font(run, size_pt=22, bold=True, color=NAVY_700, hebrew_hint=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bot Specification & Tracking Dashboard")
    set_run_font(run, size_pt=14, italic=True, color=GRAY_500)

    # Spacer
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Subtitle / metadata
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Vendor-Ready Specification")
    set_run_font(run, size_pt=13, color=TEAL_600, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version 0.1  ·  2026-05-27")
    set_run_font(run, size_pt=11, color=GRAY_500)

    # More spacer
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Confidential footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL — internal use only")
    set_run_font(run, size_pt=9, color=GRAY_500, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Drafted for the incoming bot vendor handoff")
    set_run_font(run, size_pt=9, color=GRAY_500, italic=True)

    doc.add_page_break()


def add_toc(doc):
    """Add a real Word-native TOC field that updates on open."""
    p = doc.add_paragraph()
    set_paragraph_rtl(p, True)
    run = p.add_run("תוכן עניינים")
    set_run_font(run, size_pt=24, bold=True, color=NAVY_900, hebrew_hint=True)
    p.paragraph_format.space_after = Pt(14)

    # TOC field — Word will populate on first open with F9 or "update fields"
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click → Update Field after opening, to populate the TOC."
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_begin)
    run._r.append(instr)
    run._r.append(fldChar_sep)
    run._r.append(placeholder)
    run._r.append(fldChar_end)
    set_run_font(run, size_pt=10, italic=True, color=GRAY_500)

    doc.add_page_break()


# ============================================================
# Markdown parser → docx
# ============================================================

def parse_md_to_docx(doc, md_text, base_heading_offset=2):
    """
    Convert a single markdown file into the document.

    base_heading_offset: # in MD becomes Heading <offset>. So if offset=2,
    # H1 in MD → Heading 2 in docx. ## → Heading 3. etc.
    Max heading level supported by docx is 6.
    """
    lines = md_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # --- Code block ---
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].rstrip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # consume the closing ```
            add_code_block(doc, "\n".join(code_lines))
            continue

        # --- Horizontal rule ---
        if re.match(r"^---+\s*$", stripped):
            add_hr(doc)
            i += 1
            continue

        # --- Heading ---
        h_match = re.match(r"^(#{1,6})\s+(.+?)\s*$", stripped)
        if h_match:
            level = len(h_match.group(1)) + base_heading_offset - 1
            level = min(level, 6)
            text = h_match.group(2)
            # Strip leading numeric prefix like "0.1 " or "9.2 " for cleaner headings
            text = re.sub(r"^\d+(\.\d+)*\s+", "", text)
            add_heading(doc, text, level)
            i += 1
            continue

        # --- Table ---
        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\s*\|?[\s\-:|]+\|?\s*$", lines[i + 1]):
            # collect table
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            add_md_table(doc, table_lines)
            continue

        # --- Bullet list ---
        if re.match(r"^[-*]\s+", stripped):
            list_items = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].rstrip()):
                item_text = re.sub(r"^[-*]\s+", "", lines[i].rstrip())
                # Continuation lines
                while i + 1 < len(lines) and lines[i + 1].startswith("  ") and not re.match(r"^[-*]\s+", lines[i + 1].rstrip()):
                    item_text += " " + lines[i + 1].strip()
                    i += 1
                list_items.append(item_text)
                i += 1
            add_bullet_list(doc, list_items)
            continue

        # --- Numbered list ---
        if re.match(r"^\d+\.\s+", stripped):
            list_items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].rstrip()):
                item_text = re.sub(r"^\d+\.\s+", "", lines[i].rstrip())
                list_items.append(item_text)
                i += 1
            add_numbered_list(doc, list_items)
            continue

        # --- Blockquote ---
        if stripped.startswith("> "):
            quote_lines = []
            while i < len(lines) and lines[i].rstrip().startswith(">"):
                ql = re.sub(r"^>\s?", "", lines[i].rstrip())
                quote_lines.append(ql)
                i += 1
            add_blockquote(doc, "\n".join(quote_lines))
            continue

        # --- Checkbox item (- [ ] or - [x]) ---
        cb_match = re.match(r"^-\s+\[([ xX])\]\s+(.+)$", stripped)
        if cb_match:
            checked = cb_match.group(1).lower() == "x"
            text = cb_match.group(2)
            add_checkbox_item(doc, text, checked)
            i += 1
            continue

        # --- Paragraph (default) ---
        if stripped == "":
            i += 1
            continue

        # Gather paragraph (joining wrapped lines)
        para_lines = [stripped]
        while i + 1 < len(lines) and lines[i + 1].strip() and not _is_special_line(lines[i + 1]):
            para_lines.append(lines[i + 1].strip())
            i += 1
        add_paragraph(doc, " ".join(para_lines))
        i += 1


def _is_special_line(line):
    s = line.rstrip()
    if s.startswith(("#", "```", ">")):
        return True
    if re.match(r"^---+\s*$", s):
        return True
    if re.match(r"^[-*]\s+", s) or re.match(r"^\d+\.\s+", s):
        return True
    if "|" in s and s.strip().startswith("|"):
        return True
    return False


# ============================================================
# Block builders
# ============================================================

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    if has_significant_hebrew(text):
        set_paragraph_rtl(p, True)
    add_inline_runs(p, text,
                    base_size=int(doc.styles[f"Heading {level}"].font.size.pt),
                    hebrew_hint=has_significant_hebrew(text))


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    if has_significant_hebrew(text):
        set_paragraph_rtl(p, True)
    add_inline_runs(p, text, base_size=11, hebrew_hint=has_significant_hebrew(text))


def add_code_block(doc, code):
    """
    Add a multi-line ``` block. Auto-detect whether it's:
    - Hebrew message text (the bot's outgoing copy) → RTL, sans Heebo, teal border, shaded
    - Technical code (SQL/JSON/predicate) → monospace, gray-shaded
    """
    is_hebrew_message = has_significant_hebrew(code)

    # Use a 1-cell table for the shaded background and border
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]

    # Shade
    fill = "F0FAF8" if is_hebrew_message else "F6F8FA"
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

    # For Hebrew messages, add a teal accent border on the side that's
    # visually "leading" in RTL (the right side).
    if is_hebrew_message:
        tcBorders = OxmlElement("w:tcBorders")
        for side, weight, color in [
            ("right", "24", "14B8A6"),  # leading edge of RTL block
            ("top", "4", "D4DCEE"),
            ("bottom", "4", "D4DCEE"),
            ("left", "4", "D4DCEE"),
        ]:
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), weight)
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), color)
            tcBorders.append(b)
        tcPr.append(tcBorders)

    # Add lines
    cell.paragraphs[0].text = ""
    for idx, line in enumerate(code.split("\n")):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.25 if is_hebrew_message else 1.15

        if is_hebrew_message:
            set_paragraph_rtl(p, True)
            run = p.add_run(line if line else " ")
            set_run_font(run, size_pt=11, color=NAVY_900,
                         mono=False, hebrew_hint=True)
        else:
            # Per-line Hebrew detection: even inside a mostly-English code
            # block (like an ASCII funnel chart), individual lines may contain
            # Hebrew. Consolas has no Hebrew glyphs, so for those lines use
            # Heebo to keep the characters readable.
            line_has_hebrew = any('֐' <= c <= '׿' for c in line)
            run = p.add_run(line if line else " ")
            if line_has_hebrew:
                set_run_font(run, size_pt=10, mono=False, color=CODE_FG,
                             hebrew_hint=True)
            else:
                set_run_font(run, size_pt=9.5, mono=True, color=CODE_FG)

    # Spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D4DCEE")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)


def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if has_significant_hebrew(item):
            set_paragraph_rtl(p, True)
        p.paragraph_format.space_after = Pt(2)
        add_inline_runs(p, item, base_size=11, hebrew_hint=has_significant_hebrew(item))


def add_numbered_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        if has_significant_hebrew(item):
            set_paragraph_rtl(p, True)
        p.paragraph_format.space_after = Pt(2)
        add_inline_runs(p, item, base_size=11, hebrew_hint=has_significant_hebrew(item))


def add_checkbox_item(doc, text, checked):
    p = doc.add_paragraph()
    if has_significant_hebrew(text):
        set_paragraph_rtl(p, True)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    mark = "☑ " if checked else "☐ "
    run = p.add_run(mark)
    set_run_font(run, size_pt=11, color=TEAL_600 if checked else GRAY_500)
    add_inline_runs(p, text, base_size=11, hebrew_hint=has_significant_hebrew(text))


def add_blockquote(doc, text):
    """Add a blockquote-style block: indented, left border, italic."""
    for line in text.split("\n"):
        p = doc.add_paragraph()
        if has_significant_hebrew(line):
            set_paragraph_rtl(p, True)
            add_paragraph_border(p, position="right", color="14B8A6", size=24)
            p.paragraph_format.right_indent = Cm(0.4)
        else:
            add_paragraph_border(p, position="left", color="14B8A6", size=24)
            p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(4)
        add_inline_runs(p, line, base_size=11, hebrew_hint=has_significant_hebrew(line))


def add_md_table(doc, table_lines):
    """Parse a markdown table and add it to the document."""
    # Remove leading/trailing pipes, split cells
    def parse_row(line):
        cells = line.strip().strip("|").split("|")
        return [c.strip() for c in cells]

    header = parse_row(table_lines[0])
    # table_lines[1] is the separator
    data_rows = [parse_row(l) for l in table_lines[2:] if l.strip()]

    table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    # Header row
    for col_idx, cell_text in enumerate(header):
        cell = table.rows[0].cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        if has_significant_hebrew(cell_text):
            set_paragraph_rtl(p, True)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(cell_text)
        set_run_font(run, size_pt=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                     hebrew_hint=has_significant_hebrew(cell_text))
        # Header shading
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "16243E")
        tcPr.append(shd)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for row_idx, row_data in enumerate(data_rows, start=1):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= len(header):
                continue
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            if has_significant_hebrew(cell_text):
                set_paragraph_rtl(p, True)
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(p, cell_text, base_size=10,
                            hebrew_hint=has_significant_hebrew(cell_text))
            # Alternate row shading
            if row_idx % 2 == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F4F7FC")
                tcPr.append(shd)

    # Spacer after table
    doc.add_paragraph()


# ============================================================
# Part-page (cover for each Part)
# ============================================================

def add_part_page(doc, part_hebrew, part_english):
    doc.add_page_break()

    # Vertical spacer
    for _ in range(7):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Hebrew title — uses Heading 1 style (so it appears in TOC) but
    # we override alignment to center and font size to be larger.
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rtl(p, True)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(part_hebrew)
    set_run_font(run, size_pt=32, bold=True, color=NAVY_900, hebrew_hint=True)

    # Teal underline bar
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("─────────")
    set_run_font(run, size_pt=18, color=TEAL_600)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(part_english)
    set_run_font(run, size_pt=14, italic=True, color=GRAY_500)

    doc.add_page_break()


# ============================================================
# Build
# ============================================================

def build():
    doc = setup_document()

    add_title_page(doc)
    add_toc(doc)

    for part in PLAN:
        add_part_page(doc, part["part"], part["part_en"])

        for chapter_title, filename in part["files"]:
            # Chapter heading (Heading 2 — Part is Heading 1)
            add_heading(doc, chapter_title, 2)

            # Read and parse the source file
            src_path = part["dir"] / filename
            if not src_path.exists():
                print(f"WARNING: missing {src_path}")
                continue
            md_text = src_path.read_text(encoding="utf-8")

            # Strip the file's own top-level # heading (we already added chapter title)
            md_text = re.sub(r"^#\s+.+?\n", "", md_text, count=1)

            # Source attribution footnote
            p = doc.add_paragraph()
            run = p.add_run(f"Source: {part['dir'].name}/{filename}")
            set_run_font(run, size_pt=8.5, italic=True, color=GRAY_500)
            p.paragraph_format.space_after = Pt(10)

            # Parse the rest. base_heading_offset=2 means MD ## → docx Heading 3.
            # (MD # was already stripped; chapter title used H2.)
            parse_md_to_docx(doc, md_text, base_heading_offset=2)

            # Section break between chapters (page break for cleanliness)
            doc.add_page_break()

    # Save
    doc.save(OUTPUT_PATH)
    print(f"OK - built: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
