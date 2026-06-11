"""
Build the non-technical PM spec DOCX (with full scenario texts) from the single
source markdown file PM_SPEC_full.md.

Reuses the proven RTL/font helpers from build_docx.py (same folder).

Run:
    PYTHONIOENCODING=utf-8 python build_pm_spec.py

Output:
    PM_SPEC_full.docx (in this docs/ folder)
"""

import re
import sys
from pathlib import Path

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Reuse the helpers already proven out in this workspace
from build_docx import (
    setup_document, add_toc, parse_md_to_docx, set_paragraph_rtl, set_run_font,
    NAVY_900, NAVY_700, TEAL_600, GRAY_500,
)

SRC = Path(__file__).resolve().parent.parent / "PM_SPEC_full.md"
OUTPUT_PATH = Path(__file__).resolve().parent / "PM_SPEC_full.docx"


def add_title_page(doc):
    for _ in range(6):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Road Protect")
    set_run_font(run, size_pt=38, bold=True, color=NAVY_900)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("━━━━━━━━━━")
    set_run_font(run, size_pt=14, color=TEAL_600)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rtl(p, True)
    run = p.add_run("בוט WhatsApp ולוח בקרה — אפיון מוצר")
    set_run_font(run, size_pt=24, bold=True, color=NAVY_700, hebrew_hint=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("WhatsApp Bot & Dashboard — Product Spec")
    set_run_font(run, size_pt=14, italic=True, color=GRAY_500)

    for _ in range(2):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Non-technical — user journeys, full bot copy, value prop, UI")
    set_run_font(run, size_pt=12.5, bold=True, color=TEAL_600)

    for _ in range(7):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Q2 2026  ·  Owner: Yossi  ·  internal")
    set_run_font(run, size_pt=11, color=GRAY_500)

    doc.add_page_break()


def build():
    md = SRC.read_text(encoding="utf-8")

    # The title page covers the document title, so drop the leading
    # "# Road Protect ..." H1 and the "## Full PM Spec ..." H2 to avoid a
    # duplicate cover heading. Keep everything from the first real section.
    md = re.sub(r"^#\s+Road Protect.*?\n", "", md, count=1)
    md = re.sub(r"^##\s+Full PM Spec.*?\n", "", md, count=1)

    doc = setup_document()
    add_title_page(doc)
    add_toc(doc)

    # Single file: MD "# " → docx Heading 1. offset=1.
    parse_md_to_docx(doc, md, base_heading_offset=1)

    doc.save(OUTPUT_PATH)
    print(f"OK - built: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
