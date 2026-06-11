"""
Build the Hebrew narrative vendor brief from vendor-narrative.md.

This is the prose-only, vendor-facing version. The richer technical
spec lives in spec.docx; this is the "story" doc.

Run:
    python build_vendor_brief.py

Output:
    vendor-brief.docx (in this same docs/ folder)
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# Configuration
# ============================================================

DOCS_DIR = Path(__file__).resolve().parent
SOURCE_PATH = DOCS_DIR / "vendor-narrative.md"
OUTPUT_PATH = DOCS_DIR / "vendor-brief.docx"

# Brand colors
NAVY_900 = RGBColor(0x0D, 0x18, 0x28)
NAVY_700 = RGBColor(0x22, 0x35, 0x5C)
NAVY_500 = RGBColor(0x3B, 0x5A, 0x9C)
TEAL_600 = RGBColor(0x0D, 0x94, 0x88)
TEAL_500 = RGBColor(0x14, 0xB8, 0xA6)
GRAY_700 = RGBColor(0x37, 0x41, 0x51)
GRAY_500 = RGBColor(0x6B, 0x72, 0x80)
BODY_BLACK = RGBColor(0x1A, 0x1A, 0x1A)

# Fonts
FONT_HEBREW = "Heebo"
FONT_LATIN = "Calibri"

HEBREW_RANGE = re.compile(r"[֐-׿]")


# ============================================================
# Low-level helpers
# ============================================================

def has_significant_hebrew(text: str) -> bool:
    if not text:
        return False
    letters = [c for c in text if c.isalpha()]
    if not letters:
        return False
    hebrew = [c for c in letters if HEBREW_RANGE.match(c)]
    return len(hebrew) / len(letters) >= 0.20


def set_paragraph_rtl(paragraph, rtl: bool = True):
    pPr = paragraph._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:bidi")):
        pPr.remove(existing)
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1" if rtl else "0")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT


def set_run_font(run, size_pt=12, bold=False, italic=False,
                 color=BODY_BLACK, hebrew=True):
    font = run.font
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    if color is not None:
        font.color.rgb = color
    font.name = FONT_LATIN
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_LATIN)
    rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    rFonts.set(qn("w:cs"), FONT_HEBREW)

    if hebrew:
        # CRITICAL: Mark the run itself as RTL. Without <w:rtl/> on the run,
        # Word may render Hebrew characters in LTR order even when the
        # paragraph is bidi — producing reversed-looking Hebrew.
        if rPr.find(qn("w:rtl")) is None:
            rPr.append(OxmlElement("w:rtl"))
        # Hebrew locale for proper bidi handling
        lang = rPr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            rPr.append(lang)
        lang.set(qn("w:bidi"), "he-IL")
        # Complex-script attributes for size/bold/italic
        if bold:
            if rPr.find(qn("w:bCs")) is None:
                rPr.append(OxmlElement("w:bCs"))
        if italic:
            if rPr.find(qn("w:iCs")) is None:
                rPr.append(OxmlElement("w:iCs"))
        szCs = rPr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rPr.append(szCs)
        szCs.set(qn("w:val"), str(int(size_pt * 2)))


# ============================================================
# Document setup — Hebrew-first
# ============================================================

def setup_document():
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)

    # Normal style — Hebrew first, generous line spacing, ~12pt
    style = doc.styles["Normal"]
    style.font.name = FONT_LATIN
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.55

    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style.element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_LATIN)
    rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    rFonts.set(qn("w:cs"), FONT_HEBREW)

    # Heading styles
    heading_specs = {
        1: dict(size=24, color=NAVY_900, before=32, after=14),
        2: dict(size=17, color=NAVY_700, before=20, after=10),
        3: dict(size=13.5, color=NAVY_500, before=14, after=6),
        4: dict(size=12, color=NAVY_500, before=10, after=4),
    }

    for level, spec in heading_specs.items():
        s = doc.styles[f"Heading {level}"]
        s.font.name = FONT_LATIN
        s.font.size = Pt(spec["size"])
        s.font.bold = True
        s.font.color.rgb = spec["color"]
        s.paragraph_format.space_before = Pt(spec["before"])
        s.paragraph_format.space_after = Pt(spec["after"])
        s.paragraph_format.keep_with_next = True

        rPr = s.element.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            s.element.insert(0, rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), FONT_LATIN)
        rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        rFonts.set(qn("w:cs"), FONT_HEBREW)

    return doc


# ============================================================
# Inline markdown — bold, italic, inline-code
# ============================================================

INLINE_RE = re.compile(
    r"(\*\*([^*]+?)\*\*"
    r"|\*([^*\n]+?)\*"
    r"|`([^`]+?)`)"
)


def add_inline(paragraph, text, base_size=12, hebrew=True):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            chunk = text[pos:m.start()]
            run = paragraph.add_run(chunk)
            set_run_font(run, size_pt=base_size, hebrew=hebrew)

        if m.group(2) is not None:  # bold
            run = paragraph.add_run(m.group(2))
            set_run_font(run, size_pt=base_size, bold=True,
                         color=NAVY_900, hebrew=hebrew)
        elif m.group(3) is not None:  # italic
            run = paragraph.add_run(m.group(3))
            set_run_font(run, size_pt=base_size, italic=True, hebrew=hebrew)
        elif m.group(4) is not None:  # inline code — treated as emphasized name
            run = paragraph.add_run(m.group(4))
            set_run_font(run, size_pt=base_size, color=TEAL_600,
                         italic=True, hebrew=hebrew)
        pos = m.end()

    if pos < len(text):
        chunk = text[pos:]
        if chunk:
            run = paragraph.add_run(chunk)
            set_run_font(run, size_pt=base_size, hebrew=hebrew)


# ============================================================
# Title page
# ============================================================

def add_title_page(doc):
    # Top spacer
    for _ in range(7):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Brand mark
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rtl(p, True)
    run = p.add_run("Road Protect")
    set_run_font(run, size_pt=40, bold=True, color=NAVY_900, hebrew=False)

    # Teal divider
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("━━━━━━━━━━")
    set_run_font(run, size_pt=14, color=TEAL_500, hebrew=False)

    # Hebrew title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rtl(p, True)
    run = p.add_run("אפיון הבוט ולוח הבקרה")
    set_run_font(run, size_pt=28, bold=True, color=NAVY_700, hebrew=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rtl(p, True)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("מסמך נרטיבי לוונדור")
    set_run_font(run, size_pt=15, italic=True, color=GRAY_500, hebrew=True)

    # Spacer
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Date / version
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rtl(p, True)
    run = p.add_run("גרסה 1.0  ·  מאי 2026")
    set_run_font(run, size_pt=12, color=GRAY_500, hebrew=True)

    # More spacer
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Confidential
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rtl(p, True)
    run = p.add_run("חסוי — לשימוש פנימי וצוות הוונדור בלבד")
    set_run_font(run, size_pt=10, italic=True, color=GRAY_500, hebrew=True)

    doc.add_page_break()


def add_toc(doc):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, True)
    run = p.add_run("תוכן עניינים")
    set_run_font(run, size_pt=22, bold=True, color=NAVY_900, hebrew=True)
    p.paragraph_format.space_after = Pt(16)

    # Native Word TOC field — auto-populated on first open
    p = doc.add_paragraph()
    set_paragraph_rtl(p, True)
    run = p.add_run()
    f_begin = OxmlElement("w:fldChar")
    f_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    f_sep = OxmlElement("w:fldChar")
    f_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "לחץ ימני → Update Field לאחר פתיחת המסמך כדי לרענן את תוכן העניינים."
    f_end = OxmlElement("w:fldChar")
    f_end.set(qn("w:fldCharType"), "end")
    run._r.append(f_begin)
    run._r.append(instr)
    run._r.append(f_sep)
    run._r.append(placeholder)
    run._r.append(f_end)
    set_run_font(run, size_pt=11, italic=True, color=GRAY_500, hebrew=True)

    doc.add_page_break()


# ============================================================
# Block builders
# ============================================================

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    is_heb = has_significant_hebrew(text)
    if is_heb:
        set_paragraph_rtl(p, True)
    size = int(doc.styles[f"Heading {level}"].font.size.pt)
    add_inline(p, text, base_size=size, hebrew=is_heb)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    is_heb = has_significant_hebrew(text)
    if is_heb:
        set_paragraph_rtl(p, True)
    # Hebrew text gets slightly larger size for readability
    base_size = 12 if is_heb else 11.5
    add_inline(p, text, base_size=base_size, hebrew=is_heb)


def add_message_block(doc, lines):
    """
    A blockquote in the source MD is a Hebrew bot message.
    Render as: shaded cell, teal accent border on the right (RTL leading edge),
    Heebo sans, slightly larger, with RTL alignment.
    """
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]

    tcPr = cell._tc.get_or_add_tcPr()

    # Background — very light teal/mint
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0FAF8")
    tcPr.append(shd)

    # Teal-accent right border (= RTL leading), thin border elsewhere
    tcBorders = OxmlElement("w:tcBorders")
    for side, weight, color in [
        ("right", "32", "14B8A6"),
        ("top", "4", "C8E8E2"),
        ("bottom", "4", "C8E8E2"),
        ("left", "4", "C8E8E2"),
    ]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), weight)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

    # Cell margins
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "bottom", "left", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "180")
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

    cell.paragraphs[0].text = ""
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.35
        set_paragraph_rtl(p, True)
        if line.strip():
            add_inline(p, line, base_size=12, hebrew=True)
        else:
            run = p.add_run(" ")
            set_run_font(run, size_pt=12, hebrew=True)

    # Spacer paragraph after the table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)


def add_bullet(doc, item):
    p = doc.add_paragraph(style="List Bullet")
    is_heb = has_significant_hebrew(item)
    if is_heb:
        set_paragraph_rtl(p, True)
    p.paragraph_format.space_after = Pt(4)
    add_inline(p, item, base_size=12 if is_heb else 11.5, hebrew=is_heb)


def add_section_separator(doc):
    """Treat `---` as a soft visual break (a centered diamond) between major sections."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("◆ ◆ ◆")
    set_run_font(run, size_pt=11, color=TEAL_500, hebrew=False)


# ============================================================
# Markdown parser
# ============================================================

def parse(doc, md_text):
    lines = md_text.splitlines()
    i = 0
    last_heading_was_h1 = False

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Section separator
        if re.match(r"^---+\s*$", stripped):
            # Skip — page breaks happen on H1 (handled below)
            i += 1
            continue

        # Heading
        h = re.match(r"^(#{1,4})\s+(.+?)\s*$", stripped)
        if h:
            level = len(h.group(1))
            text = h.group(2)
            # Strip prefixes like "## 2. " or "# 1."
            text = re.sub(r"^(\d+(\.\d+)*)\.\s+", "", text)
            # H1 starts on a new page (except the first one)
            if level == 1 and last_heading_was_h1 is False and i > 0:
                # Add a page break before H1 if it's not the first content
                # (skip for first H1 which immediately follows the TOC)
                pass  # We'll let page break be controlled at H1 start below
            if level == 1:
                # Always start H1 on a new page except for the very first
                doc.add_page_break()
            add_heading(doc, text, level)
            last_heading_was_h1 = (level == 1)
            i += 1
            continue

        last_heading_was_h1 = False

        # Blockquote — a bot message
        if stripped.startswith("> "):
            msg_lines = []
            while i < len(lines) and lines[i].rstrip().startswith(">"):
                msg_lines.append(re.sub(r"^>\s?", "", lines[i].rstrip()))
                i += 1
            add_message_block(doc, msg_lines)
            continue

        # Bullet
        if re.match(r"^-\s+", stripped):
            while i < len(lines) and re.match(r"^-\s+", lines[i].rstrip()):
                item = re.sub(r"^-\s+", "", lines[i].rstrip())
                add_bullet(doc, item)
                i += 1
            continue

        # Blank line
        if stripped == "":
            i += 1
            continue

        # Paragraph — gather wrapped lines
        para_lines = [stripped]
        while (
            i + 1 < len(lines)
            and lines[i + 1].strip()
            and not lines[i + 1].rstrip().startswith(("#", "> ", "-", "---"))
        ):
            para_lines.append(lines[i + 1].rstrip())
            i += 1
        add_paragraph(doc, " ".join(para_lines))
        i += 1


# ============================================================
# Build
# ============================================================

def build():
    md_text = SOURCE_PATH.read_text(encoding="utf-8")

    # Strip the top H1 (title + subtitle lines) — we'll build a custom title page
    # The source starts with:
    # # Road Protect — אפיון הבוט ולוח הבקרה
    # ## מסמך אפיון לוונדור · גרסה 1.0 · מאי 2026
    md_text = re.sub(r"^#\s+Road Protect.+?\n", "", md_text, count=1)
    md_text = re.sub(r"^##\s+מסמך אפיון.+?\n", "", md_text, count=1)
    # Strip the leading --- after that
    md_text = re.sub(r"^---+\s*\n", "", md_text, count=1)

    doc = setup_document()
    add_title_page(doc)
    add_toc(doc)

    parse(doc, md_text)

    doc.save(OUTPUT_PATH)
    print(f"OK - built: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
